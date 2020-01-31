#!/usr/bin/env python
# coding: utf-8

# # <center> Shut Up & Sit Down Podscast Analysis </center>
# *We are not kidding, that's their name.*

# ![](https://www.shutupandsitdown.com/wp-content/uploads/2019/01/shut-up-logo@2x.png)

# ## Description and goal of the analysis

# Shut Up & Sit Down is a group of people that review and create content centered around board games. They write blog posts, create YouTube videos and record a podcast. We'll be focusing on the last part in this analysis. 
# 
# During the podcast they talk about a wide variety of games, from many genres. We want to see if there are some interesting patterns in their work. For this, we decided to combine their podcast with information from the biggest board game database in the world - Board Game Geek. 

# ## Some word about the data

# We'll be using two datasets:
# 
# ### Transcripts of the Shut Up & Sit Down Podcast
# Source: https://drive.google.com/drive/folders/1SXN8c3WHWLX0HSDUjwpOSIWfwGRxRlBB <br>
# Provided by the community. They follow a strict formatting guide that helps greatly when we are importing and processing the data. 
# 
# We are fans of SU&SD and we watch their videos. However, we have not listened to the podcast yet. This is actually a good factor, because it allows us to minimize hidden bias. 
# 
# ### Board Game Geek Database 
# Source: https://www.kaggle.com/jvanelteren/boardgamegeek-reviews <br>
# Consists of 17 000 games scrapped from the BGG website in 2019. Includes many variables, such as a rank, rating average, year of publishing and many more. 

# <hr>

# ## Setup
# ### Import the libraries
# A pretty standard combinaation. We'll be using nltk for text processing.

# In[146]:


import pandas as pd
import numpy as np
import string
import difflib # Fuzzy matching
import csv 
from datetime import datetime

# Plotting
import matplotlib.pyplot as plt 
import seaborn as sns
from wordcloud import WordCloud

import re
from docx import Document
import nltk
from itertools import chain
import heapq

from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize 
from nltk.sentiment.vader import SentimentIntensityAnalyzer


# ### Set display settings and colors
# This way we can get around the limit of iPython and see all columns in the data. We'll also import some nice colors for plotting later (inspired by Airbnb).

# In[2]:


pd.set_option('display.max_columns', 500)
pd.set_option('display.width', 1000)
pd.set_option('display.float_format', lambda x: '%.2f' %x)
pink = "#FF5A5F"
orange = "#FFB400"
blueGreen = "#007A87"
flesh = "#FFAA91"
purple = "#7B0051"


# ### Center the plots using CSS
# In our opinion plots should be centered by default. 

# In[3]:


from IPython.core.display import HTML

HTML("""
<style>
.output_png {
    display: table-cell;
    text-align: center;
    margin:auto;
    }
.prompt 
    display:none;
}  
</style>
""")


# ### Downloads for text preparation
# We need to download punctuation, stop words and a lexicon for sentiment analysis. 

# In[ ]:


nltk.download('punkt')
nltk.download('stopwords')
nltk.download('wordnet')
nltk.download('averaged_perceptron_tagger')
nltk.download('vader_lexicon')


# <hr>

# # Data import and transformation

# ### Import the Word documents
# We changed the names of the documents to make them easier to import ;)

# In[4]:


document_list = []
for i in range(9):
    document = Document('data/raw/Podcast_9{}.docx'.format(i))
    document_list.append(document)


# ### Import the Board Game Geek data
# 17 000 games and 56 columns. Not all will be usefull, but we'll deal with that later. 

# In[5]:


game_data = pd.read_csv('data\games_detailed_info.csv', low_memory=False)


# In[7]:


game_data.describe()


# ### Define a function to extract headings
# Just a loop through all paragraphs.

# In[63]:


def iter_headings(paragraphs):
    for paragraph in paragraphs:
        if paragraph.style.name.startswith('Heading'):
            yield paragraph


# ### Proces the documents
# Each section in the documents begins with a header with the name of the game, then each paragraph is started with info about the speaker. 

# In[64]:


results = pd.DataFrame()

for z in document_list:
    
    # Create an empty list for all headings
    heading_list = []
    
    # Extraction headings from this document
    for heading in iter_headings(z.paragraphs):
        heading_list.append(heading.text)

    # Remove the intro and outro
    heading_list = heading_list[1:-1]
    
    text_lines = []

    for i,para in enumerate(z.paragraphs):
        if para.text in heading_list:
            text_lines.append(i)

    #extracting text lines for games only (without intro and outro)       
    gamesText = []
    for k,j in enumerate(text_lines):
        for y, parar in enumerate(z.paragraphs):
            if y>=text_lines[k]:
                gamesText.append(parar.text)

    #creating list mapping where we get the in-text position of heading (number of line)            
    text_lines2 = [text_lines[n]-text_lines[0] for n in range(1,len(text_lines))]
    text_lines2.insert(0,0)
    
    #Extracting text with respect to headings: list of dialogs related to specific game
    extracted_text = []

    for d,w in zip(text_lines2,text_lines2[1:]):
        text = gamesText[d+1:w]
        extracted_text.append(text) 
    
    for game, text in zip(heading_list, extracted_text):
        for sentence in text:
            results = results.append({'game' : game, 'text': sentence}, ignore_index = True)


# ### Make changes to the text
# We need to extract the person speaking and put that information into a different column. We'll also drop some paragraphs as they contain useless data. We'll also remove the timestamp from the headers.

# In[65]:


# Remove the timestamp
results['game'] = results['game'].map(lambda x: str(x)[:-11])

# Split the text into person and speech
new = results["text"].str.split(":", n = 1, expand = True) 
  
# Extract the person speaking
results["person"] = new[0] 
  
# And what they are saying
results["text"] = new[1] 

# Drop if the formating was off
results.dropna(inplace = True)

# Just a preacusion
results[results['person'].map(len) < 12]

# Remove the part after their name
results["person"] = results["person"].str.split(" ", n = 1, expand = True)[0]

# Remove "This" and "And"
results = results[(results['person'] != 'This') & (results['person'] != 'And')]

# One of the games is wrong
results.replace('Twilight Imperium stream plug', 'Twilight Imperium', inplace = True)

# Create a list of games
games = results.game.unique()


# As we can see, each paragraph is a row in the Data Frame.

# In[66]:


results


# ### Create per-game text table
# We'll also create a second table with all the text about a single game put into one row. This will allow us to do easy per-game analysis. 

# In[67]:


grouped_results = dict()

for index, row in results.iterrows():
    game = row[0]
    text = row[1]
    if game not in grouped_results:
        grouped_results[game] = text
    else:
        grouped_results[game] = grouped_results[game] + ' ' + text
        
grouped_results = pd.DataFrame(list(grouped_results.items()), columns=['game', 'text'])


# In[68]:


grouped_results


# <hr>

# # Text transformation

# ### Tokenization
# We'll hold the processed words in a different column. First, we split them by " ". 

# In[69]:


grouped_results['formated_text'] = grouped_results['text'].map(lambda x: x.split())
results['formated_text'] = results['text'].map(lambda x: x.split())


# ### Remove punctuation and stop words
# Next, we'll remove stopwords, lowercase the words and remove any punctuation on a per letter basis.

# In[70]:


def remove_words(text):
    keep = []
    stop_words = list(stopwords.words('english')) + ['I', 'i', 'a', 'A', "it's"]
    punctuation = list(string.punctuation) + ["''", "_", "“", "’", "…", "[", "-", "”", "–"]
    
    for word in text:
        word = word.lower()
        if word not in stop_words:
            clean_word = ''
            for letter in word:
                if letter not in punctuation:
                    clean_word += letter.lower()
                    
            if clean_word != '': 
                keep.append(clean_word)      
    return keep

grouped_results['formated_text'] = grouped_results['formated_text'].map(remove_words)
results['formated_text'] = results['formated_text'].map(remove_words)


# <hr>

# # Merge the podcasts with BGG 

# It is time to merge the datasets. Since the names will not be exactly equal, we'll use fuzzy matching and pick the closest one. 

# In[71]:


known_matches = dict()

def match_game(game):
    if game not in known_matches.keys():
        match = difflib.get_close_matches(game, game_data['primary'], n = 1)
        if not match:
            known_matches[game] = ''
            return ''
        else:
            known_matches[game] = match[0]
            return match[0]
    else: 
        return known_matches[game]

grouped_results['primary'] = grouped_results['game'].map(match_game)
results['primary'] = results['game'].map(match_game)


# ### Merge the datasets
# Good results - we dropped only 10 games out of 66. 

# In[72]:


grouped_results = grouped_results.merge(game_data, on = 'primary', how = 'inner')
results = results.merge(game_data, on = 'primary', how = 'inner')


# ### Remove useless columns
# Not every column is usefull. The ranks below have only a few observations. Some of the other variables are purely technical in nature, and thous of no use. 

# In[74]:


grouped_results.drop(columns = [
    'Unnamed: 0',
    'Accessory Rank',
    'Amiga Rank', 
    'Arcade Rank', 
    'Atari ST Rank', 
    'Commodore 64 Rank', 
    'Customizable Rank', 
    'RPG Item Rank',
    'alternate',
    'boardgameartist',
    'boardgamecategory',
    'boardgamecompilation',
    'boardgamedesigner',
    'boardgameexpansion',
    'boardgamefamily', 
    'boardgameimplementation',
    'boardgameintegration',
    'boardgamemechanic', 
    'boardgamepublisher',
    'id', 
    'image', 
    'suggested_language_dependence', 
    'suggested_num_players',
    'suggested_playerage', 
    'thumbnail',
    'type', 
], inplace = True)

results.drop(columns = [
    'Unnamed: 0',
    'Accessory Rank',
    'Amiga Rank', 
    'Arcade Rank', 
    'Atari ST Rank', 
    'Commodore 64 Rank', 
    'Customizable Rank', 
    'RPG Item Rank',
    'alternate',
    'boardgameartist',
    'boardgamecategory',
    'boardgamecompilation',
    'boardgamedesigner',
    'boardgameexpansion',
    'boardgamefamily', 
    'boardgameimplementation',
    'boardgameintegration',
    'boardgamemechanic', 
    'boardgamepublisher',
    'id', 
    'image', 
    'suggested_language_dependence', 
    'suggested_num_players',
    'suggested_playerage', 
    'thumbnail',
    'type', 
], inplace = True)

grouped_results = grouped_results.astype({'Board Game Rank': 'int64'})
results = results.astype({'Board Game Rank': 'int64'})


# <hr>

# # Calculate the sentiment 
# The sentiment analysis will be performed using the **VADER model**. We could explain that it is a good fit for our type of data, but we'll be honest. We picked it because of the awesome name.
# 
# We also create a new variable - **sentiment ratio**. It is calculated by dividing the positive sentiment by the negative sentiment. We'll test later if this metric will be of any use. 

# In[75]:


def analize_sentiment(words):
    sid = SentimentIntensityAnalyzer()
    
    text = ''
    for word in words: text = text + ' ' + word

    return sid.polarity_scores(text)

grouped_results['positive'] = grouped_results['formated_text'].map(lambda x: analize_sentiment(x)['pos'])
grouped_results['neutral'] = grouped_results['formated_text'].map(lambda x: analize_sentiment(x)['neu'])
grouped_results['negative'] = grouped_results['formated_text'].map(lambda x: analize_sentiment(x)['neg'])
grouped_results['sentiment_ratio'] = grouped_results['positive'] / grouped_results['negative']

results['positive'] = results['formated_text'].map(lambda x: analize_sentiment(x)['pos'])
results['neutral'] = results['formated_text'].map(lambda x: analize_sentiment(x)['neu'])
results['negative'] = results['formated_text'].map(lambda x: analize_sentiment(x)['neg'])
results['sentiment_ratio'] = results['positive'] / results['negative']


# <hr>

# # Calculate word frequency for all reviews
# Now we'll also calculate the frequency of words in all reviews. We'll use the data without the stopwords etc.

# In[76]:


frequency = {}

for index, row in grouped_results.iterrows():
    text = row[2]
    
    for word in text:
        if word not in frequency.keys():
            frequency[word] = 1
        else:
            frequency[word] += 1

# Convert to a dataframe
frequency = pd.DataFrame(list(frequency.items()), columns=['word', 'frequency'])

# Sort by frequency
frequency.sort_values('frequency', ascending = False, inplace = True)

frequency.reset_index(drop = True, inplace = True)


# In[77]:


frequency


# <hr>

# # Explore the data

# ## Does the Zipf law apply here?
# Looks like Zipf was onto something. The distribution indeed follows the path described by the french stenographer. 

# In[78]:


fig, ax = plt.subplots(figsize=(15,10))
ax.plot(frequency['frequency'], color = purple)
plt.title('Word frequency')
plt.ylabel('Frequency')
plt.show()


# In[79]:


fig, ax = plt.subplots(figsize=(15,10))
ax.plot(frequency['frequency'][0:100], color = purple)
plt.title('Word frequency for the 100 most frequent words')
plt.ylabel('Frequency')
plt.show()


# In[80]:


plt.figure(figsize=(15,10))
plt.barh('word', 'frequency', data = frequency[0:20], color = "#FFB400")
plt.gca().invert_yaxis()
plt.title('Most popular words')
plt.ylabel('Word')
plt.xlabel('Frequency')
full+text = plt.show()


# We can also see a nice wordcloud to get a better sense of the data. *Game* is obvious, but *thing* is an interesting effect of the dynamic no-script nature of the podcast. 

# In[157]:


full_text = ''

for index, row in grouped_results.iterrows():
    text = row[2]
    
    for word in text:
        full_text = full_text + ' ' + word

wordcloud = WordCloud(width=1920, height=1080, background_color = 'white').generate(full_text)
plt.figure(figsize=(15,10))
plt.imshow(wordcloud)
plt.axis("off")
plt.show()


# ## Are the variables correlated? 
# There are a significant number of correlated variables in the dataset. Sadly, it seems that there are not many correlations between databases. 
# 
# Of note is an interesting pattern - positive sentimend in negatively correlated with the neutral sentiment, but a similar correlation between negative and neutral is much weaker. 
# 
# Positive and negative sentiment also has only a weak negative correlation. 

# In[42]:


fig, ax = plt.subplots(figsize=(20,20))
corr = grouped_results[['positive',
                        'neutral', 
                        'negative', 
                        'Board Game Rank',
                        'average',
                        'averageweight', 
                        'bayesaverage',
                        'maxplayers', 
                        'maxplaytime',
                        'minage', 
                        'minplayers',
                        'minplaytime',
                        'numcomments',
                        'numweights', 
                        'owned', 
                        'playingtime',
                        'stddev',
                        'trading', 
                        'usersrated', 
                        'wanting',
                        'wishing',
                        'yearpublished',
                        'sentiment_ratio']].corr()
ax = sns.heatmap(
    corr, 
    vmin=-1, vmax=1, center=0,
    cmap="PuOr",
    square=True
)
ax.set_xticklabels(
    ax.get_xticklabels(),
    rotation=45,
    horizontalalignment='right'
);
plt.title('Correlation matrix of the numerical variables')
plt.show()


# If we run the analisys on the paragraphs and not on the text, all correlations between datasets dissapear.

# In[81]:


fig, ax = plt.subplots(figsize=(20,20))
corr = results[['positive',
                        'neutral', 
                        'negative', 
                        'Board Game Rank',
                        'average',
                        'averageweight', 
                        'bayesaverage',
                        'maxplayers', 
                        'maxplaytime',
                        'minage', 
                        'minplayers',
                        'minplaytime',
                        'numcomments',
                        'numweights', 
                        'owned', 
                        'playingtime',
                        'stddev',
                        'trading', 
                        'usersrated', 
                        'wanting',
                        'wishing',
                        'yearpublished',
                        'sentiment_ratio']].corr()
ax = sns.heatmap(
    corr, 
    vmin=-1, vmax=1, center=0,
    cmap="PuOr",
    square=True
)
ax.set_xticklabels(
    ax.get_xticklabels(),
    rotation=45,
    horizontalalignment='right'
);
plt.title('Correlation matrix of the numerical variables')
plt.show()


# ## What are the patterns of sentiment in the reviews?
# It looks like the sentiment for all reviews hover around means. There are no significant outliers, all games follow the patternm. 
# 
# For neutral the mean would be ~ 65%, for positive ~ 25% and for negative ~ 10%. This means that for the most part, the reviewers focus on the games. When they voice their opinions it is more than 2x likely that they the opinion will be positive. 
# 
# This is an important information, because this may indicate that most games choosen for the podcast are good, or at least, are liked by the participants. 

# In[45]:


fig, ax = plt.subplots(figsize=(15,10))
ax.plot('positive', data = grouped_results, color = purple)
ax.plot('neutral', data = grouped_results, color = orange)
ax.plot('negative', data = grouped_results, color = blueGreen)
handles, labels = ax.get_legend_handles_labels()
ax.legend(handles, labels)
plt.title('Distribution of sentiment')
plt.ylabel('% of text')
plt.xlabel('Game')
plt.show()


# The same plot as above, but in a way that is probably more clear. 

# In[112]:


ax = grouped_results[['negative', 'neutral', 'positive']].plot(kind = 'bar',
                                                               stacked = True,
                                                               figsize = (15,10),
                                                               width = 1,
                                                               color = [blueGreen, orange, purple],
                                                               title = 'Distribution of sentiment')


# Sadly the x axis cannot be removed here, so it looks weird.
# 
# Interesting pattern - for each paragraph the variance of sentiment is much bigger. Some are very positive, some very neutral. Some event very negative, but these are rare.

# In[115]:


ax = results[['negative', 'neutral', 'positive']].plot(kind = 'bar',
                                                       stacked = True,
                                                       figsize = (25,20),
                                                       width = 1,
                                                       color = [blueGreen, orange, purple],
                                                       title = 'Distribution of sentiment per paragraph')


# ## What about the sentiment ratio?
# This measure provides interesting insight. It looks like for some games there is nearly 10 times more positive content than negative!
# 
# Note that no games reach negative numbers, and most stay above 2 (that is 2x more positive than negative).

# In[84]:


fig, ax = plt.subplots(figsize=(15,10))
ax.plot('sentiment_ratio', data = grouped_results, color = purple)
plt.title('Sentiment ratio for each game')
plt.ylabel('Sentiment ratio')
plt.xlabel('Game')
plt.show()


# ## Let's confirm the spread of sentiment 
# The values on the plot axis confirm that the reviews are mostly positive. Interestingly, there are no other patterns here.

# In[122]:


fig, ax = plt.subplots(figsize=(15,10))
ax.scatter('positive', 'negative', c = 'neutral', data = grouped_results, cmap="Wistia")
plt.title('Positive vs negative sentiment (color is neutral)')
plt.ylabel('Negative sentiment')
plt.xlabel('Positive sentiment')
plt.show()


# When we look at the paragraph data, interesting conclusions can be drawn. Many paragraphs have 0% of either positive or negative sentiment. The minimal amount of sentiment seems to be about 5% for this data. 

# In[121]:


fig, ax = plt.subplots(figsize=(15,10))
ax.scatter('positive', 'negative', c = 'neutral', data = results, cmap="Wistia")
plt.title('Positive vs negative sentiment (color is neutral)')
plt.ylabel('Negative sentiment')
plt.xlabel('Positive sentiment')
plt.show()


# ## Are the ratings correlated with the sentiment?
# It seems that there is no correlation between the ratings and the sentiment of the review. The games here represent the typical scale of averages on BGG (6 to 8). 

# In[134]:


fig, ax = plt.subplots(figsize=(15,10))
ax.scatter('positive', 'average', color = orange, data = grouped_results)
ax.scatter('negative', 'average', color = purple, data = grouped_results)
ax.scatter('neutral', 'average', color = blueGreen, data = grouped_results)
plt.show()


# The same result is repeated on the paragraph data. No patterns here.

# In[139]:


fig, ax = plt.subplots(figsize=(15,10))
ax.scatter('positive', 'average', color = orange, data = results, alpha = 0.1)
ax.scatter('negative', 'average', color = purple, data = results, alpha = 0.1)
ax.scatter('neutral', 'average', color = blueGreen, data = results, alpha = 0.1)
plt.show()


# In[141]:


fig, ax = plt.subplots(figsize=(15,10))
ax.scatter('positive', 'average', color = orange, data = results, alpha = 0.5)
plt.show()


# ## What about the sentiment ratio?
# Sadly, it seems that there is no correlation here either. 

# In[142]:


plt.figure(figsize=(15,10))
plt.scatter('sentiment_ratio', 'average', color = orange, data = grouped_results)
plt.show()


# Same for the paragraph data.

# In[144]:


plt.figure(figsize=(15,10))
plt.scatter('sentiment_ratio', 'average', color = orange, data = results, alpha = 0.5)
plt.show()


# ## Bonus: sentiment per person

# *"Quinn person" and "audience" can be ignored.*
# 
# Some very interesting patterns here. Quinns and Matt are the two typical hosts and they seem to have simmilar proportions of sentiment. Kylie (an intern) seems to have much stronger opinions. He is much more likely to offer positive opinions, less likely to offer neutral ones and a bit more likely to offer negative ones.

# In[178]:


plt.figure(figsize=(15,10))
sns.violinplot(x = "person", 
               y = "positive",
               split = True, 
               inner = "quart",
               data = results,
              palette = sns.cubehelix_palette())
plt.title('Positive sentiment per person')
sns.despine(left = True)


# In[179]:


plt.figure(figsize=(15,10))
sns.violinplot(x = "person", 
               y = "neutral",
               split = True, 
               inner = "quart",
               data = results,
              palette = sns.cubehelix_palette())
plt.title('Positive sentiment per person')
sns.despine(left = True)


# In[180]:


plt.figure(figsize=(15,10))
sns.violinplot(x = "person", 
               y = "negative",
               split = True, 
               inner = "quart",
               data = results,
              palette = sns.cubehelix_palette())
plt.title('Positive sentiment per person')
sns.despine(left = True)


# # Conclusion
# We were able to draw some interesting conclusions from this analysis:
# 
# The data follows the Zipf law of frequency. This is true for all reviews. 
# 
# The most popular words are related to the nature of the podcast: game, thing, people, think. 
# 
# The people in the podcast seem to mostly talk about the games in a neutral way. When they voice their opinions, it is typically in a positive way. They do sometimes criticise the games, but out of 55 games analized, none had more negative sentiment than positive. 
# 
# There seems to be no connection between the sentiment of the review and the average score on Board Game Geek. There may be many reasons for this, but we think this may have to do with the baias of the games - the persenters mostly talk about games they like. 
# 
# It seems that two main hosts: Quinns and Matt have simmilar sentiment proportions. This may be one of the reasons why they were able to keep working together for many years. The ocassional quest - Kylie the Intern is more likely to voice positive opinions than Quinns and Matt. 
